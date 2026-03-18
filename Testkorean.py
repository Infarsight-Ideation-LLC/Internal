import os
import json
from datetime import datetime
from playwright.sync_api import sync_playwright
from docx import Document
from openai import AzureOpenAI
 
 
URL = "https://englishdart.fss.or.kr/dsbh001/main.do?rcpNo=20260310901403"
 
 
client = AzureOpenAI(
    api_key="4Ng4vnx2KNiWNF9PTdG79Df5kWcV3S9Ci62WQvdwh1Dpr4FphNFQJQQJ99CAACHYHv6XJ3w3AAAAACOG2jHs",
    api_version="2025-01-01-preview",
    azure_endpoint="https://rs-ai-internal-projects-poc-dev.cognitiveservices.azure.com"
)
 
 
# ==============================
# SCRAPE WEBSITE
# ==============================
 
def scrape_text():
 
    with sync_playwright() as p:
 
        browser = p.chromium.launch(headless=False)
        page = browser.new_page()
 
        print("Opening page...")
 
        page.goto(URL, timeout=60000)
 
        page.wait_for_load_state("domcontentloaded")
        page.wait_for_load_state("networkidle")
 
        page.wait_for_timeout(3000)
 
        frames = page.frames
 
        full_text = ""
 
        for frame in frames:
            try:
                text = frame.locator("body").inner_text(timeout=5000)
                full_text += text + "\n"
            except:
                pass
 
        browser.close()
 
        return full_text
 
 
# ==============================
# OPENAI EXTRACTION
# ==============================
 
def extract_data(text):
 
    prompt = f"""
 
Read the disclosure text and return:
 
1️⃣ Summary in English
2️⃣ Extract these parameters
 
Return JSON only.
 
Parameters to extract:
 
Project / Asset Name
Country / City
Sector / Sub-sector
Investment / Deal Value
Currency
Sponsors / Investors
Lenders / Banks
Project Status / Stage
Source / Publication Date
 
If a value is missing write "Not Found".
 
TEXT:
{text}
 
 
Output JSON format:
 
{{
"summary":"English summary here",
"parameters":{{
"Project / Asset Name":"",
"Country / City":"",
"Sector / Sub-sector":"",
"Investment / Deal Value":"",
"Currency":"",
"Sponsors / Investors":"",
"Lenders / Banks":"",
"Project Status / Stage":"",
"Source / Publication Date":""
}}
}}
 
"""
 
    response = client.chat.completions.create(
        model="IFS_Internal_POC_Dev",
        messages=[
            {"role": "system", "content": "You extract structured financial project data."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.2
    )
 
    result = response.choices[0].message.content
 
    return json.loads(result)
 
 
# ==============================
# CREATE WORD DOCUMENT
# ==============================
 
def create_word(data):
 
    summary = data["summary"]
    parameters = data["parameters"]
 
    onedrive_path = os.getenv("OneDriveCommercial") or os.getenv("OneDrive")
 
    if not onedrive_path:
        raise Exception("OneDrive folder not found")
 
    output_folder = os.path.join(onedrive_path, "IJ Global Extracted File")
 
    os.makedirs(output_folder, exist_ok=True)
 
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
 
    file_name = f"DART_Summary_{timestamp}.docx"
 
    file_path = os.path.join(output_folder, file_name)
 
    doc = Document()
 
    doc.add_heading("Korean DART Disclosure Summary", level=0)
 
    doc.add_paragraph(f"Generated on: {datetime.now()}")
 
    doc.add_heading("Source URL", level=1)
    doc.add_paragraph(URL)
 
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(summary)
 
    doc.add_paragraph("")
 
    # ============================
    # PARAMETERS TABLE
    # ============================
 
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
 
    headers = table.rows[0].cells
    headers[0].text = "Parameters"
    headers[1].text = "Value"
 
    # Bold header
    for cell in headers:
        run = cell.paragraphs[0].runs[0]
        run.bold = True
 
    for key, value in parameters.items():
 
        row = table.add_row().cells
 
        row[0].text = key
        row[1].text = value
 
    doc.save(file_path)
 
    # Generate JSON metadata
    json_file_path = file_path.replace(".docx", ".json")
    metadata = {
        "file_name": file_name,
        "project_name": parameters.get("Project / Asset Name", "Not Found"),
        "country": "South Korea",
        "region": "APAC",
        "industry_type": parameters.get("Sector / Sub-sector", "Not Found"),
        "generated_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "file_size_kb": round(os.path.getsize(file_path) / 1024, 1),
        "source_url": URL,
        "website": "Korea Dart"
    }
    
    with open(json_file_path, "w") as json_file:
        json.dump(metadata, json_file, indent=2)
 
    print("Word file saved to OneDrive:")
    print(file_path)
    print("Metadata saved to:")
    print(json_file_path)
 
 
# ==============================
# MAIN
# ==============================
 
def run():
 
    print("Extracting website text...")
 
    text = scrape_text()
 
    print("Sending to Azure OpenAI...")
 
    data = extract_data(text)
 
    print("\nSummary:\n")
    print(data["summary"])
 
    print("\nCreating Word file...")
 
    create_word(data)
 
 
if __name__ == "__main__":
    run()