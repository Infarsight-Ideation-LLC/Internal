import json
import pprint
from datetime import datetime, timedelta
from playwright.sync_api import sync_playwright
from docx import Document
import os
import re
 
from NewsIntentOpenAI import detect_news_intent
from ParametersExtract import extract_project_details
 
 
def run_scraper():
 
    print("Opening sites.json...")
 
    with open("config/sites.json", "r") as f:
        data = json.load(f)
 
    sites = data["sites"]
 
    all_projects_text = []
 
    with sync_playwright() as p:
 
        browser = p.chromium.launch(headless=False)
        page = browser.new_page()
 
        # increase timeout stability
        page.set_default_timeout(60000)
 
        for site in sites:
 
            if site["valid"] != "Yes":
                print("Skipping site:", site["siteName"])
                continue
 
            print("\nOpening site:", site["siteName"])
 
            page.goto(site["siteURL"], timeout=60000)
 
            page.wait_for_load_state("domcontentloaded")
            page.wait_for_load_state("networkidle")
            page.wait_for_timeout(2000)
 
            industry = site["industryType"]
 
            print("Selecting Industry:", industry)
 
            page.select_option("#dropdown_3", label=industry)
 
            days = site["NoOfDays"]
 
            date_from = (datetime.today() - timedelta(days=days)).strftime("%Y-%m-%d")
            date_to = datetime.today().strftime("%Y-%m-%d")
 
            print("Setting date range:", date_from, "to", date_to)
 
            page.evaluate(
                """([selector, value]) => {
                    const el = document.querySelector(selector);
                    el.value = value;
                    el.dispatchEvent(new Event('input', { bubbles: true }));
                    el.dispatchEvent(new Event('change', { bubbles: true }));
                }""",
                ["#date_4", date_from]
            )
 
            page.evaluate(
                """([selector, value]) => {
                    const el = document.querySelector(selector);
                    el.value = value;
                    el.dispatchEvent(new Event('input', { bubbles: true }));
                    el.dispatchEvent(new Event('change', { bubbles: true }));
                }""",
                ["#date_5", date_to]
            )
 
            print("Clicking Apply Filter")
 
            page.click("button:has-text('Apply Filter')")
 
            page.wait_for_selector("table tbody tr", timeout=60000)
 
            rows = page.query_selector_all("table tbody tr")
 
            print("Total records found:", len(rows))
 
            for i in range(len(rows)):
 
                rows = page.query_selector_all("table tbody tr")
 
                link = rows[i].query_selector("td:first-child a")
 
                if not link:
                    continue
 
                epbc_number = link.inner_text()
 
                print("\nOpening record:", epbc_number)
 
                # CLICK RECORD
                link.click()

                # Wait for a key section or fallback to body (handles AJAX/dynamic navigation)
                try:
                    page.wait_for_selector("text=Project description", timeout=20000)
                except Exception:
                    page.wait_for_selector("body", timeout=10000)

                # extra wait for dynamic content
                page.wait_for_timeout(3000)
 
                # scroll to bottom (lazy loading fix)
                page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
                page.wait_for_timeout(1500)
 
                # extract full page text
                text = page.locator("body").inner_text(timeout=30000)
 
                all_projects_text.append({
                    "epbc_number": epbc_number,
                    "text": text
                })
 
                # go back
                page.go_back()
 
                page.wait_for_load_state("domcontentloaded")
                page.wait_for_load_state("networkidle")
                page.wait_for_timeout(2000)
 
                page.click("button:has-text('Apply Filter')")
 
                page.wait_for_selector("table tbody tr")
 
        browser.close()
 
    print("\n===== FULL ARRAY DATA =====\n")
    pprint.pprint(all_projects_text)
 
    print("\nStarting OpenAI Processing...\n")
 
    for project in all_projects_text:
 
        print("Processing:", project["epbc_number"])
 
        intent = detect_news_intent(project["text"])
        print("AI Intent:", intent)
 
        structured_data = extract_project_details(project["text"])
 
        print("\nExtracted JSON:\n")
        print(structured_data)
 
        output_dir = r"C:\Users\Admin\OneDrive - Infarsight Ideation LLP\IJ Global Extracted File"
        os.makedirs(output_dir, exist_ok=True)
 
        if isinstance(structured_data, str):
            try:
                structured_data = json.loads(structured_data)
            except Exception:
                structured_data = {}
 
        project_name = structured_data.get("Project Name", "Unknown Project")
        project_summary = structured_data.get("Project Summary", "No summary available.")
 
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
 
        epbc_safe = project["epbc_number"].replace("/", "_")
 
        project_name_safe = re.sub(r'[\\/*?:"<>|]', "", project_name)
        project_name_safe = project_name_safe.replace(" ", "")
 
        file_name = f"{epbc_safe}_{project_name_safe}_{timestamp}.docx"
 
        file_path = os.path.join(output_dir, file_name)
 
        doc = Document()
 
        doc.add_heading(project_name, level=0)
 
        gen_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        doc.add_paragraph(f"Generated on: {gen_time}")
 
        doc.add_heading("Project Summary", level=1)
        doc.add_paragraph(project_summary)
 
        doc.add_paragraph("")
 
        table = doc.add_table(rows=1, cols=2)
 
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Field'
        hdr_cells[1].text = 'Value'
 
        for key, value in structured_data.items():
 
            if key in ["Project Name", "Project Summary"]:
                continue
 
            row_cells = table.add_row().cells
            row_cells[0].text = str(key)
            row_cells[1].text = str(value)
 
        doc.save(file_path)
 
        print(f"Saved Word document: {file_path}")
        print("\n----------------------------------\n")
 
 
if __name__ == "__main__":
    run_scraper()